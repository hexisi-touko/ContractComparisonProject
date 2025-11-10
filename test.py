# -*- coding: utf-8 -*-
import sys
import os
import re
import shutil
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QWidget, QFileDialog, QMessageBox,
                             QVBoxLayout, QHBoxLayout, QLabel, QPushButton)
from PyQt6.QtWidgets import QApplication, QWidget, QFileDialog, QMessageBox
import mammoth
from difflib import SequenceMatcher
from bs4 import BeautifulSoup
from ui.optimized_compare import Ui_Form
from PyQt6.QtCore import Qt  # 注意：PyQt6 中是小写的 qt（区分大小写）
from docx import Document
from docx.shared import RGBColor

class CompareApp(QWidget, Ui_Form):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("文件对比工具")

        # 初始化历史文件存储目录
        self.history_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "history_files")
        if not os.path.exists(self.history_dir):
            os.makedirs(self.history_dir)

        # 绑定按钮事件
        self.importOriginalFileButton.clicked.connect(self.load_original_file)
        self.importCompareFileButton.clicked.connect(self.load_compare_file)
        self.compareButton.clicked.connect(self.compare_files)
        self.historyButton.clicked.connect(self.show_history_page)
        self.exportButton.clicked.connect(self.export_highlighted_file)
        # 存储路径与HTML内容
        self.original_file_path = None
        self.compare_file_path = None
        self.original_html = None
        self.compare_html = None
        self.original_text_blocks = []
        self.compare_text_blocks = []
        self.highlighted_html = None  # 保存标红后的HTML结果
        # 历史页面实例（作为子窗口）
        self.history_page = None

        # 模拟 Word 样式
        self.word_css = """
        <style>
            body {
                font-family: 'SimSun', '宋体', serif;
                font-size: 12pt;
                line-height: 1.8;
                color: #000000;
                background-color: #ffffff;
                max-width: 21cm;
                margin: 0 auto;
                padding: 2.5cm 2cm;
            }
            .contract-main-title {
                text-align: center;
                font-size: 16pt;
                font-weight: bold;
                margin: 0 0 40pt 0;
                text-indent: 0;
            }
            .party-info {
                text-align: left;
                margin: 15pt 0;
                text-indent: 0;
            }
            .contract-preface {
                text-align: justify;
                text-indent: 2em;
                margin: 20pt 0;
            }
            .clause-level1 {
                font-weight: bold;
                font-size: 13pt;
                margin: 25pt 0 10pt 0;
                text-indent: 0;
            }
            .clause-level2 {
                font-weight: bold;
                text-indent: 2em;
                margin: 15pt 0 5pt 0;
            }
            p {
                text-align: justify;
                text-indent: 2em;
                margin: 8pt 0;
            }
            ul, ol {
                margin: 5pt 0 5pt 4em;
                padding-left: 0;
            }
            li {
                text-align: justify;
                margin: 6pt 0;
                text-indent: 0;
            }
            .signature-area {
                margin-top: 60pt;
                text-indent: 0;
            }
            .signature-item {
                margin: 20pt 0;
                text-indent: 0;
            }
            .sign-date {
                margin-top: 30pt;
                text-indent: 0;
            }
            .diff-highlight {
                background-color: #ffcccc;
                color: #ff0000;
                font-weight: bold;
                padding: 0 2px;
            }
            .diff-delete {
                color: #ff0000;
                text-decoration: line-through;
                font-weight: bold;
                padding: 0 2px;
            }
            /* 在word_css中增加样式 */
            .level-change-only { 
                color: #ff9900; 
                background-color: #fff8e1;  /* 浅黄背景区分纯层级变化 */
            }
        </style>
        """

    # --------------------------------------------------------
    # 从 HTML 提取文本块
    # --------------------------------------------------------
    def extract_text_blocks(self, html_content):
        soup = BeautifulSoup(html_content, 'html.parser')
        text_blocks = []
        # 扩展需要提取的标签类型（覆盖合同常见元素）
        target_tags = ['p', 'li', 'h1', 'h2', 'h3', 'td', 'th']  # 新增标题、表格单元格标签

        for tag in target_tags:
            for element in soup.find_all(tag):
                # 提取文本并过滤纯空白/无效内容
                text = element.get_text(strip=True)
                if not text:
                    continue  # 跳过空文本块

                # 提取合同层级信息（基于CSS类名，如 clause-level1 对应一级条款）
                level = 0
                classes = element.get('class', [])
                for cls in classes:
                    if cls.startswith('clause-level'):
                        # 从类名中提取层级数字（如 'clause-level2' → 2）
                        level = int(cls.split('-')[-1])
                        break

                # 提取合同特有标识（如条款号、角色名）
                identifier = None
                # 匹配条款号（如"第1条"、"1.1"、"一、"等）
                clause_pattern = re.compile(r'^(第?\d+[条款项]|[\d.]+|[\u4e00-\u9fa5]+、)')
                if clause_pattern.match(text):
                    identifier = clause_pattern.match(text).group()

                # 存储块信息（包含标签类型、层级、标识）
                text_blocks.append({
                    'text': text,
                    'tag': tag,
                    'level': level,  # 用于结构化匹配
                    'identifier': identifier,  # 用于锚点定位（如"第3条"）
                    'original_element': element  # 保留原始元素，便于后续还原格式
                })

        return text_blocks

    def compare_files(self):
        if not self.original_file_path or not self.compare_file_path:
            QMessageBox.warning(self, "警告", "请先导入原文件和对比文件！")
            return
        if not self.original_text_blocks or not self.compare_text_blocks:
            QMessageBox.warning(self, "警告", "文件内容解析失败，请重新导入！")
            return

        try:
            # 1. 基于条款标识和层级的智能匹配（核心优化）
            matched_pairs = self.match_blocks_by_structure()
            if not matched_pairs:
                QMessageBox.warning(self, "提示", "未找到可匹配的条款结构，将使用默认顺序对比")
                # 退回到原始顺序对比逻辑
                matched_pairs = [(i, i) for i in
                                 range(min(len(self.original_text_blocks), len(self.compare_text_blocks)))]
                extra_compare_indices = list(range(len(matched_pairs), len(self.compare_text_blocks)))
            else:
                # 提取未匹配的新增条款
                extra_compare_indices = [j for j in range(len(self.compare_text_blocks))
                                         if j not in [pair[1] for pair in matched_pairs]]

            # 2. 初始化对比文档
            soup = BeautifulSoup(self.compare_html, 'html.parser')
            compare_nodes = soup.find_all(['p', 'li', 'h1', 'h2', 'h3', 'td', 'th'])  # 扩展支持的标签
            diff_count = 0

            # 3. 对比已匹配的条款块
            for orig_idx, comp_idx in matched_pairs:
                orig_block = self.original_text_blocks[orig_idx]
                comp_block = self.compare_text_blocks[comp_idx]
                node = compare_nodes[comp_idx]

                # 跳过空文本块
                if not orig_block['text'] or not comp_block['text']:
                    continue

                # 4. 针对合同关键信息的增强对比
                highlighted_html = self.highlight_differences(orig_block['text'], comp_block['text'])

                # 5. 标记条款层级变化（如一级条款变成二级条款）
                # 优化层级变化判断：文本相同则仅标记不计数，文本不同则正常计数
                if orig_block.get('level') != comp_block.get('level'):
                    highlighted_html = f'<span class="level-change">[层级变化] {highlighted_html}</span>'
                    # 只有当文本内容不同时，才计入差异计数
                    if highlighted_html != comp_block['text']:
                        diff_count += 1

                if highlighted_html != comp_block['text']:
                    diff_count += 1
                    node.string = ''
                    node.append(BeautifulSoup(highlighted_html, 'html.parser'))

            # 6. 标记新增条款（合同中新增的条款单独标注来源）
            for comp_idx in extra_compare_indices:
                node = compare_nodes[comp_idx]
                extra_text = node.get_text(strip=True)
                if extra_text:
                    diff_count += 1
                    # 新增条款标记中加入原文件位置提示（如"新增于原文件第X条后"）
                    insert_pos = self.get_insert_position(comp_idx)
                    highlight_html = f'<span class="diff-highlight">[新增条款{insert_pos}] {extra_text}</span>'
                    node.string = ''
                    node.append(BeautifulSoup(highlight_html, 'html.parser'))

            # 7. 标记原文件有但对比文件缺失的条款
            missing_indices = [i for i in range(len(self.original_text_blocks))
                               if i not in [pair[0] for pair in matched_pairs]]
            if missing_indices:
                # 在对比文档末尾添加缺失条款汇总
                missing_section = soup.new_tag('div')
                missing_section['class'] = 'missing-clauses'
                missing_section.append(BeautifulSoup('<p><strong>原文件缺失条款：</strong></p>', 'html.parser'))

                for orig_idx in missing_indices:
                    orig_text = self.original_text_blocks[orig_idx]['text']
                    missing_p = soup.new_tag('p')
                    # 确保 BeautifulSoup 解析结果正确
                    span_soup = BeautifulSoup(f'<span class="diff-delete">[缺失] {orig_text}</span>', 'html.parser')
                    missing_p.append(span_soup)
                    missing_section.append(missing_p)
                    diff_count += 1

                # 关键修复：检查并确保 body 标签存在
                if soup.body is None:
                    body = soup.new_tag('body')
                    soup.append(body)
                soup.body.append(missing_section)

            # 8. 生成最终HTML
            highlighted_full_html = f"""
            <!DOCTYPE html><html><head>
            <meta charset="UTF-8">{self.word_css}
            <style>
                .level-change {{ color: #ff9900; }}  /* 层级变化标记为橙色 */
                .missing-clauses {{ margin-top: 20pt; padding: 10pt; border: 1px solid #ff0000; }}
            </style>
            </head><body>
            {str(soup)}</body></html>
            """

            self.highlighted_html = highlighted_full_html
            self.webEngineCompareView.setHtml(highlighted_full_html)
            QMessageBox.information(self, "完成", f"文件对比完成！共发现 {diff_count} 处差异（含条款新增/缺失/层级变化）。")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"文件对比失败：\n{e}")

    # 新增：基于条款结构的匹配方法（需配合优化后的extract_text_blocks使用）
    def match_blocks_by_structure(self):
        """通过条款标识（如第1条）和层级匹配对应文本块，解决顺序变动问题"""
        matched = []
        comp_matched = set()  # 记录已匹配的对比文件索引

        # 优先通过条款标识匹配（如"第1条"必须匹配）
        for orig_idx, orig_block in enumerate(self.original_text_blocks):
            orig_id = orig_block.get('identifier')
            if not orig_id:
                continue
            # 在对比文件中找相同标识的条款
            for comp_idx, comp_block in enumerate(self.compare_text_blocks):
                if comp_idx in comp_matched:
                    continue
                # 优先通过条款标识匹配时，增加层级相似性判断
                if comp_block.get('identifier') == orig_id:
                    # 若层级差异过大（如相差>1级），降低匹配优先级
                    level_diff = abs(comp_block.get('level', 0) - orig_block.get('level', 0))
                    if level_diff <= 1:  # 允许相邻层级的微小差异
                        matched.append((orig_idx, comp_idx))
                        comp_matched.add(comp_idx)
                        break

        # 剩余未匹配项按层级+文本相似度匹配
        for orig_idx, orig_block in enumerate(self.original_text_blocks):
            if orig_idx in [p[0] for p in matched]:
                continue  # 跳过已匹配项
            orig_level = orig_block.get('level', 0)
            orig_text = orig_block['text']
            # 找同层级且文本相似度>0.7的条款
            for comp_idx, comp_block in enumerate(self.compare_text_blocks):
                if comp_idx in comp_matched:
                    continue
                if comp_block.get('level', 0) != orig_level:
                    continue
                # 计算文本相似度
                similarity = SequenceMatcher(None, orig_text, comp_block['text']).ratio()
                if similarity > 0.7:
                    matched.append((orig_idx, comp_idx))
                    comp_matched.add(comp_idx)
                    break

        return matched

    # 新增：计算新增条款在原文件中的插入位置提示
    def get_insert_position(self, comp_idx):
        """判断新增条款在原文件中的相对位置（如"第3条后"）"""
        comp_block = self.compare_text_blocks[comp_idx]
        comp_level = comp_block.get('level', 0)
        # 找到原文件中同层级的最后一个条款
        last_orig_idx = -1
        for i, block in enumerate(self.original_text_blocks):
            if block.get('level', 0) == comp_level:
                last_orig_idx = i
        if last_orig_idx == -1:
            return ""
        orig_id = self.original_text_blocks[last_orig_idx].get('identifier', f"第{last_orig_idx + 1}项")
        return f"（位于原文件{orig_id}后）"

    def highlight_differences(self, original_text, compare_text):
        """增强版差异标红：优化中文分词、忽略无关空格、支持标点符号精确对比"""
        if original_text == compare_text:
            return compare_text

        # 1. 预处理：统一空格和换行符（合同中常因格式产生无关空格差异）
        def preprocess(text):
            # 合并连续空格为单个（保留一个空格避免语义变化）
            text = re.sub(r'\s+', ' ', text)
            # 去除首尾空格（合同条款首尾空格通常无意义）
            return text.strip()

        orig_processed = preprocess(original_text)
        comp_processed = preprocess(compare_text)

        # 预处理后仍相同则直接返回
        if orig_processed == comp_processed:
            return compare_text

        # 2. 针对中文优化的序列匹配（使用字符级比对，避免英文分词逻辑干扰）
        matcher = SequenceMatcher(None, orig_processed, comp_processed)
        result = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            orig_segment = orig_processed[i1:i2]
            comp_segment = comp_processed[j1:j2]

            if tag == 'equal':
                # 保留原始文本的空格格式（仅替换差异部分，非差异部分保持原样）
                # 从原始对比文本中截取对应片段（而非预处理后的）
                # 计算原始文本中对应位置（处理空格差异导致的索引偏移）
                # 这里简化处理：直接使用对比文本的原始片段（适合大部分场景）
                result.append(compare_text[j1:j2] if j1 < j2 else '')

            elif tag == 'insert':
                # 新增内容标红，添加"新增"提示（合同审核中需明确标识新增项）
                result.append(f'<span class="diff-highlight">[新增]{comp_segment}</span>')

            elif tag == 'delete':
                # 删除内容标红+删除线，添加"删除"提示
                result.append(f'<span class="diff-delete">[删除]{orig_segment}</span>')

            elif tag == 'replace':
                # 替换内容同时显示删除和新增部分，用"替换为"连接
                result.append(
                    f'<span class="diff-delete">[删除]{orig_segment}</span>'
                    f'<span class="diff-highlight">[替换为]{comp_segment}</span>'
                )

        # 3. 后处理：修复可能的标签嵌套问题（避免HTML解析错误）
        highlighted = ''.join(result)
        # 移除空标签（如连续删除/新增导致的空span）
        highlighted = re.sub(r'<span class="[^"]+"></span>', '', highlighted)
        return highlighted


    def load_original_file(self, file_path=None):
        """导入原文件 (.docx)，并在左侧展示区显示，支持从历史记录加载"""
        # 如果没有传入路径，则打开文件选择对话框
        if not file_path:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "选择原文件",
                "",
                "Word 文件 (*.docx)"
            )

        if not file_path:
            return  # 用户取消选择

        try:
            # 备份文件到历史文件夹（仅当不是从历史记录加载时）
            if not file_path.startswith(self.history_dir):
                # 生成带时间戳的文件名，避免重复
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = os.path.basename(file_path)
                name, ext = os.path.splitext(filename)
                history_filename = f"{name}_{timestamp}{ext}"
                history_path = os.path.join(self.history_dir, history_filename)

                # 复制文件到历史文件夹
                shutil.copy2(file_path, history_path)
                self.original_file_path = history_path
            else:
                self.original_file_path = file_path

            # 读取 docx 并转为 HTML（后续代码保持不变）
            with open(self.original_file_path, "rb") as docx_file:
                # 样式映射代码保持不变
                style_map = """
                p[style-name='标题 1'] => p.contract-main-title
                p[style-name='正文'] => p.contract-preface
                p[style-name='标题 2'] => p.clause-level1
                p[style-name='标题 3'] => p.clause-level2
                p[style-name='普通段落'] => p.party-info
                p[style-name='签名区'] => p.signature-area
                p[style-name='签名项'] => p.signature-item
                p[style-name='日期'] => p.sign-date
                """
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map,
                    # convert_image=mammoth.images.img_element

                )
                html_content = result.value

            # 组合CSS和HTML内容，使样式生效
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                {self.word_css}
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            # 加载到左侧展示区（webEngineOriginView）
            self.webEngineOriginView.setHtml(full_html)
            self.original_html = html_content
            self.original_text_blocks = self.extract_text_blocks(html_content)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法显示 Word 文件内容：\n{e}")

    def load_compare_file(self):
        """导入对比文件 (.docx)，并在右侧展示区显示"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择对比文件",
            "",
            "Word 文件 (*.docx)"
        )

        if not file_path:
            return  # 用户取消选择

        try:
            # 读取对比文件并转换为HTML（复用原文件的样式映射）
            with open(file_path, "rb") as docx_file:
                style_map = """
                p[style-name='标题 1'] => p.contract-main-title
                p[style-name='正文'] => p.contract-preface
                p[style-name='标题 2'] => p.clause-level1
                p[style-name='标题 3'] => p.clause-level2
                p[style-name='普通段落'] => p.party-info
                p[style-name='签名区'] => p.signature-area
                p[style-name='签名项'] => p.signature-item
                p[style-name='日期'] => p.sign-date
                """
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map,
                    # convert_image=mammoth.images.img_element

                )
                html_content = result.value

            # 组合CSS和HTML内容（保持与原文件格式一致）
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                {self.word_css}
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            # 加载到右侧展示区（假设控件名为 webEngineCompareView，需与UI文件一致）
            self.webEngineCompareView.setHtml(full_html)
            # 保存对比文件路径（供后续对比功能使用）
            self.compare_file_path = file_path
            self.compare_html = html_content
            self.compare_text_blocks = self.extract_text_blocks(html_content)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法显示对比文件内容：\n{e}")

    def handle_image(self,image):
        """处理图片转换，出错时返回空标签避免崩溃"""
        try:
            # 尝试转换图片为HTML标签
            return mammoth.images.img_element(image)
        except Exception:
            # 出错时返回空标签（或提示文字）
            return ("<p>[无法显示的图片]</p>",)  # 必须返回可迭代对象（如元组）

    # 在CompareApp类中添加显示历史页面的方法
    def show_history_page(self):
        """显示历史页面，大小与主界面一致并覆盖主界面"""
        # 创建历史页面（主窗口为父窗口）
        self.history_page = HistoryPage(
            parent=self,  # 关键：设置主窗口为父窗口
            history_dir=self.history_dir,
            callback=self.load_original_file
        )
        # 设置历史页面大小与主界面完全一致
        self.history_page.setGeometry(self.rect())  # 关键：同步大小和位置
        # 显示历史页面（会覆盖主界面）
        self.history_page.show()

    # 重写窗口大小改变事件：同步历史页面大小
    def resizeEvent(self, event):
        if self.history_page and self.history_page.isVisible():
            self.history_page.setGeometry(self.rect())  # 主窗口 resize 时，历史页面同步
        super().resizeEvent(event)

    # --------------------------------------------------------
    # 导出为带标红的 .docx 副本（删除部分带删除线）
    # --------------------------------------------------------
    def export_highlighted_file(self):
        if not self.highlighted_html:
            QMessageBox.warning(self, "警告", "请先完成文件对比再导出！")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word 文件副本", "对比结果.docx", "Word 文件 (*.docx)"
        )
        if not file_path:
            return

        try:
            soup = BeautifulSoup(self.highlighted_html, 'html.parser')
            doc = Document()

            for node in soup.find_all(['p', 'li']):
                para = doc.add_paragraph()
                for elem in node.contents:
                    if isinstance(elem, str):
                        run = para.add_run(elem)
                    elif elem.name == 'span' and 'diff-highlight' in elem.get('class', []):
                        run = para.add_run(elem.get_text())
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.bold = True
                    elif elem.name == 'span' and 'diff-delete' in elem.get('class', []):
                        run = para.add_run(elem.get_text())
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.strike = True  # ✅ 添加删除线
                        run.bold = True
                    else:
                        run = para.add_run(elem.get_text())

            doc.save(file_path)
            QMessageBox.information(self, "成功", f"已导出副本文件：\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：\n{e}")

# 添加历史记录页面类
class HistoryPage(QWidget):
    def __init__(self, parent=None, history_dir=None, callback=None):
        super().__init__(parent)
        self.history_dir = history_dir
        self.callback = callback  # 用于回调显示选中的历史文件

        # 1. 先设置窗口属性（在初始化UI前设置）
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.setStyleSheet("background-color: white;")  # 确保背景不透明

        # 2. 只调用一次 init_ui()（关键修复：避免重复初始化覆盖布局）
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("历史文件查询")
        layout = QVBoxLayout(self)

        # 步骤1：清空布局（确保无残留元素）
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # 步骤2：设置拉伸因子（顺序：标题0、列表1、按钮0）
        layout.setStretch(0, 0)  # 标题不拉伸
        layout.setStretch(1, 1)  # 列表占满剩余空间
        layout.setStretch(2, 0)  # 按钮不拉伸

        # 1. 标题（索引0）- 修改这里
        title_label = QLabel("历史文件列表")
        # 关键修改：通过样式表减少上下边距，限制最小高度
        title_label.setStyleSheet("""
            font: 700 14pt 'Microsoft YaHei UI';
            margin-top: 5px;    /* 顶部边距 */
            margin-bottom: 0px; /* 底部边距 */
            padding-top: 2px;   /* 内边距顶部 */
            padding-bottom: 2px; /* 内边距底部 */
            min-height: 20px;   /* 最小高度（仅够显示文字） */
            max-height: 25px;   /* 最大高度限制 */
        """)
        # 设置固定高度策略
        title_label.setFixedHeight(25)  # 强制固定高度

        layout.addWidget(title_label)

        # 2. 历史文件列表（索引1）
        self.file_list_widget = QWidget()
        self.file_list_layout = QVBoxLayout(self.file_list_widget)
        # 列表容器添加少量内边距，避免内容贴边
        self.file_list_layout.setContentsMargins(5, 5, 5, 5)  # 左、上、右、下内边距
        self.file_list_layout.setStretch(0, 1)  # 确保列表内容拉伸
        layout.addWidget(self.file_list_widget)
        self.load_history_files()

        # 3. 返回按钮（索引2）
        back_btn = QPushButton("返回主界面")
        back_btn.clicked.connect(self.close)
        layout.addWidget(back_btn)

        # 强制应用布局
        self.setLayout(layout)
        layout.update()

    def load_history_files(self):
        """加载历史文件并显示（保持不变）"""
        # 清空现有列表
        while self.file_list_layout.count():
            item = self.file_list_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # 获取并排序历史文件（按修改时间倒序）
        if os.path.exists(self.history_dir):
            files = [f for f in os.listdir(self.history_dir) if f.endswith(".docx")]
            files.sort(key=lambda x: os.path.getmtime(os.path.join(self.history_dir, x)), reverse=True)

            for file in files:
                file_path = os.path.join(self.history_dir, file)
                mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")

                # 创建文件信息和查看按钮
                file_widget = QWidget()
                file_layout = QHBoxLayout(file_widget)

                file_info = QLabel(f"{file}  ({mod_time})")
                view_btn = QPushButton("点击查看")
                view_btn.clicked.connect(lambda checked, path=file_path: self.view_file(path))

                file_layout.addWidget(file_info)
                file_layout.addWidget(view_btn)
                self.file_list_layout.addWidget(file_widget)

    def view_file(self, file_path):
        """回调主页面显示选中的历史文件（保持不变）"""
        if self.callback:
            self.callback(file_path)
        self.close()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = CompareApp()
    window.show()
    sys.exit(app.exec())